from io import BytesIO
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


class ReportTemplateAdapter:
    """按照老师提供的达成度分析 Word 样稿生成正式报告。"""

    NAME = "report_template_adapter"
    VERSION = "public-template-v1"
    DEFAULT_TEMPLATE_PATH = Path("sample_data/report_template.docx")

    @classmethod
    def build_document(cls, course, semester, class_scope, context, template_path=None):
        summary = context["summary"]
        chapter_four = summary["chapter_four"]
        chapter_five_ready = context.get("chapter_five_ready", False)

        template_document = cls._new_document(template_path)
        if cls._is_teacher_report_template(template_document) and cls._template_matches_course(template_document, course):
            cls._apply_to_existing_teacher_template(
                template_document,
                course,
                semester,
                class_scope,
                context,
            )
            return template_document

        document = cls._new_blank_document(template_path)
        cls._configure_document_styles(document)

        cls._build_cover(document, course, semester, class_scope)
        cls._start_body_section(document)
        cls._build_basic_info(document, course, semester, class_scope, summary["student_count"])
        cls._build_requirement_mapping(document, course)
        cls._build_assessment_mapping(document, course)
        cls._build_chapter_four(document, chapter_four)
        if chapter_five_ready:
            cls._build_chapter_five(document, context)
        cls._build_other_section(document, chapter_five_ready)
        cls._build_audit_section(document, chapter_five_ready)
        return document

    @classmethod
    def _new_document(cls, template_path=None):
        path = Path(template_path or cls.DEFAULT_TEMPLATE_PATH)
        return Document(str(path)) if path.exists() else Document()

    @classmethod
    def _new_blank_document(cls, template_path=None):
        document = cls._new_document(template_path)
        body = document._body._element
        for child in list(body):
            if child.tag != qn("w:sectPr"):
                body.remove(child)
        return document

    @staticmethod
    def _cell_text(cell):
        return cell.text.replace("\n", "").strip()

    @classmethod
    def _is_teacher_report_template(cls, document):
        title_found = any("课程目标达成度分析及持续改进报告" in paragraph.text for paragraph in document.paragraphs)
        enough_tables = len(document.tables) >= 5
        return title_found and enough_tables

    @classmethod
    def _template_course_identity(cls, document):
        if not document.tables:
            return "", ""
        table = document.tables[0]
        if len(table.rows) < 1 or len(table.columns) < 4:
            return "", ""
        return cls._cell_text(table.cell(0, 1)), cls._cell_text(table.cell(0, 3))

    @classmethod
    def _template_matches_course(cls, document, course):
        code, name = cls._template_course_identity(document)
        return bool(code and name and code == (course.code or "").strip() and name == (course.name or "").strip())

    @staticmethod
    def _configure_document_styles(document):
        if document.sections:
            section = document.sections[0]
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)

        normal_style = document.styles["Normal"]
        normal_style.font.name = "Calibri"
        normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        normal_style.font.size = Pt(10.5)

    @staticmethod
    def _start_body_section(document):
        document.add_page_break()
        section = document.add_section(WD_SECTION.CONTINUOUS)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)

    @staticmethod
    def _set_run_font(run, font_name="仿宋", size=16, bold=False):
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        run.font.size = Pt(size)
        run.font.bold = bold

    @classmethod
    def _add_text_paragraph(
        cls,
        document,
        text="",
        *,
        font_name="仿宋",
        size=16,
        bold=False,
        alignment=None,
        line_spacing=1.25,
        space_before=0,
        space_after=0,
        left_indent=None,
    ):
        paragraph = document.add_paragraph()
        if alignment is not None:
            paragraph.alignment = alignment
        if left_indent is not None:
            paragraph.paragraph_format.left_indent = Cm(left_indent)
        paragraph.paragraph_format.line_spacing = line_spacing
        paragraph.paragraph_format.space_before = Pt(space_before)
        paragraph.paragraph_format.space_after = Pt(space_after)
        run = paragraph.add_run(str(text))
        cls._set_run_font(run, font_name, size, bold)
        return paragraph

    @classmethod
    def _add_section_title(cls, document, text):
        return cls._add_text_paragraph(
            document,
            text,
            font_name="黑体",
            size=16,
            bold=False,
            line_spacing=1.25,
            space_before=5,
            space_after=5,
        )

    @classmethod
    def _chart_output_dir(cls):
        chart_dir = Path(__file__).resolve().parents[2] / "tmp" / "report_charts"
        chart_dir.mkdir(parents=True, exist_ok=True)
        return chart_dir

    @classmethod
    def _asset_output_dir(cls):
        asset_dir = Path(__file__).resolve().parents[2] / "tmp" / "report_assets"
        asset_dir.mkdir(parents=True, exist_ok=True)
        return asset_dir

    @classmethod
    def _cover_logo_paths(cls):
        asset_dir = cls._asset_output_dir()
        seal_path = asset_dir / "cover_seal.jpg"
        wordmark_path = asset_dir / "cover_wordmark.jpg"
        if seal_path.exists() and wordmark_path.exists():
            return seal_path, wordmark_path

        template_path = cls.DEFAULT_TEMPLATE_PATH
        if not template_path.exists():
            return None, None

        seal_blob = None
        wordmark_blob = None
        template_document = Document(str(template_path))
        for rel in template_document.part.rels.values():
            if "image" not in rel.reltype:
                continue
            blob = rel.target_part.blob
            try:
                with Image.open(BytesIO(blob)) as image:
                    width, height = image.size
            except OSError:
                continue
            ratio = width / height if height else 0
            if seal_blob is None and width < 650 and 0.8 <= ratio <= 1.2:
                seal_blob = blob
            elif wordmark_blob is None and width < 800 and 1.5 <= ratio <= 2.5:
                wordmark_blob = blob

        if seal_blob:
            seal_path.write_bytes(seal_blob)
        if wordmark_blob:
            wordmark_path.write_bytes(wordmark_blob)
        return (seal_path if seal_path.exists() else None, wordmark_path if wordmark_path.exists() else None)

    @staticmethod
    def _chart_font(size=22):
        font_candidates = [
            "/System/Library/Fonts/STHeiti Medium.ttc",
            "/System/Library/Fonts/STHeiti Light.ttc",
            "/System/Library/Fonts/Supplemental/Songti.ttc",
            "/System/Library/Fonts/PingFang.ttc",
            "/Library/Fonts/Arial Unicode.ttf",
        ]
        for font_path in font_candidates:
            if not Path(font_path).exists():
                continue
            try:
                return ImageFont.truetype(font_path, size)
            except OSError:
                continue
        return ImageFont.load_default()

    @staticmethod
    def _text_size(draw, text, font):
        box = draw.textbbox((0, 0), str(text), font=font)
        return box[2] - box[0], box[3] - box[1]

    @classmethod
    def _draw_centered_text(cls, draw, center_x, center_y, text, font, fill="#222222"):
        width, height = cls._text_size(draw, text, font)
        draw.text((center_x - width / 2, center_y - height / 2), str(text), font=font, fill=fill)

    @classmethod
    def _draw_rotated_text(cls, image, text, font, x, y, fill="#333333"):
        scratch = Image.new("RGBA", (220, 60), (255, 255, 255, 0))
        scratch_draw = ImageDraw.Draw(scratch)
        text_width, text_height = cls._text_size(scratch_draw, text, font)
        scratch = Image.new("RGBA", (text_width + 12, text_height + 12), (255, 255, 255, 0))
        scratch_draw = ImageDraw.Draw(scratch)
        scratch_draw.text((6, 6), str(text), font=font, fill=fill)
        rotated = scratch.rotate(90, expand=True)
        image.paste(rotated, (int(x), int(y)), rotated)

    @staticmethod
    def _chart_y(value, top, bottom):
        value = max(0.0, min(100.0, float(value or 0)))
        return bottom - (value / 100) * (bottom - top)

    @staticmethod
    def _as_percent(value):
        if value is None:
            return 0.0
        value = float(value or 0)
        return value * 100 if 0 < value <= 1.5 else value

    @classmethod
    def _qualitative_attainment_map(cls, summary):
        rows = summary.get("chapter_four", {}).get("qualitative_rows") or []
        return {
            row.get("objective_title"): cls._as_percent(row.get("attainment"))
            for row in rows
            if row.get("objective_title") and row.get("objective_title") != "课程总目标计算"
        }

    @classmethod
    def _objective_chart_attainment(cls, objective, qualitative_map):
        value = cls._as_percent(objective.get("quantitative_attainment_percent") or objective.get("quantitative_attainment"))
        fallback = qualitative_map.get(objective.get("objective_title"))
        if fallback and (value <= 0 or value > 110):
            return fallback
        return value

    @staticmethod
    def _values_from_qualitative_counts(objective):
        counts = objective.get("qualitative_counts") or {}
        values = []
        for label, score in (("优", 90), ("良", 80), ("中", 70), ("差", 60)):
            values.extend([score] * int(counts.get(label, 0) or 0))
        return values

    @classmethod
    def _objective_chart_values(cls, objective):
        values = []
        for item in objective.get("student_rate_percents", []):
            try:
                value = cls._as_percent(item)
            except (TypeError, ValueError):
                continue
            values.append(max(0.0, min(100.0, value)))
        if values:
            return values

        fallback_values = cls._values_from_qualitative_counts(objective)
        if fallback_values:
            return fallback_values
        return values

    @staticmethod
    def _objective_label(objective, fallback_index):
        title = str(objective.get("objective_title") or "").replace(" ", "")
        match = re.search(r"课程目标(\d+)", title)
        return f"课程目标{match.group(1)}" if match else f"课程目标{fallback_index}"

    @classmethod
    def _draw_axes(cls, draw, *, left, top, right, bottom, font):
        grid_color = "#d8d8d8"
        axis_color = "#888888"
        for tick in range(0, 101, 10):
            y = cls._chart_y(tick, top, bottom)
            draw.line((left, y, right, y), fill=grid_color, width=1)
            label = str(tick)
            label_width, label_height = cls._text_size(draw, label, font)
            draw.text((left - label_width - 12, y - label_height / 2), label, font=font, fill="#555555")
        draw.line((left, bottom, right, bottom), fill=axis_color, width=2)
        draw.line((left, top, left, bottom), fill=axis_color, width=2)

    @classmethod
    def _save_overall_attainment_chart(cls, summary, output_path):
        objectives = summary.get("objective_results", [])
        if not objectives:
            return None

        width, height = 1400, 560
        image = Image.new("RGB", (width, height), "white")
        draw = ImageDraw.Draw(image)
        title_font = cls._chart_font(34)
        label_font = cls._chart_font(22)
        small_font = cls._chart_font(18)

        cls._draw_centered_text(draw, width / 2, 38, "课程目标达成情况分析", title_font)
        left, top, right, bottom = 115, 92, 1125, 420
        cls._draw_axes(draw, left=left, top=top, right=right, bottom=bottom, font=small_font)

        expected_value = float(getattr(summary.get("course"), "expected_value", 0.65) or 0.65) * 100
        qualitative_map = cls._qualitative_attainment_map(summary)
        group_width = (right - left) / max(len(objectives), 1)
        bar_width = min(58, max(22, group_width * 0.22))
        colors = {"expected": "#4aa3df", "attainment": "#ff7600"}

        for index, objective in enumerate(objectives, start=1):
            center_x = left + group_width * (index - 0.5)
            attainment_value = cls._objective_chart_attainment(objective, qualitative_map)
            bars = [
                (expected_value, center_x - bar_width - 10, colors["expected"]),
                (attainment_value, center_x + 10, colors["attainment"]),
            ]
            for value, x0, color in bars:
                y0 = cls._chart_y(value, top, bottom)
                draw.rectangle((x0, y0, x0 + bar_width, bottom), fill=color)
            cls._draw_centered_text(draw, center_x, bottom + 36, cls._objective_label(objective, index), label_font)

        cls._draw_centered_text(draw, (left + right) / 2, height - 32, "课程目标", label_font)
        cls._draw_rotated_text(image, "达成度", label_font, 26, (top + bottom) / 2 - 42)

        legend_x, legend_y = 1170, 215
        draw.rectangle((legend_x, legend_y, legend_x + 22, legend_y + 14), fill=colors["expected"])
        draw.text((legend_x + 34, legend_y - 6), "期望值", font=label_font, fill="#333333")
        draw.rectangle((legend_x, legend_y + 38, legend_x + 22, legend_y + 52), fill=colors["attainment"])
        draw.text((legend_x + 34, legend_y + 32), "达成度值", font=label_font, fill="#333333")

        image.save(output_path)
        return output_path

    @classmethod
    def _save_objective_attainment_chart(cls, objective, output_path, fallback_index):
        values = cls._objective_chart_values(objective)
        if not values:
            return None

        width, height = 1400, 590
        image = Image.new("RGB", (width, height), "white")
        draw = ImageDraw.Draw(image)
        title_font = cls._chart_font(32)
        label_font = cls._chart_font(21)
        small_font = cls._chart_font(17)

        label = cls._objective_label(objective, fallback_index)
        cls._draw_centered_text(draw, width / 2, 36, f"{label}达成度", title_font)
        left, top, right, bottom = 105, 86, 1115, 430
        cls._draw_axes(draw, left=left, top=top, right=right, bottom=bottom, font=small_font)

        average = sum(values) / len(values)
        expected = cls._as_percent(objective.get("attainment_snapshot", {}).get("expected_value") or 0)
        average_y = cls._chart_y(average, top, bottom)
        expected_y = cls._chart_y(expected, top, bottom)
        draw.line((left, average_y, right, average_y), fill="#3f9bd8", width=3)
        draw.line((left, expected_y, right, expected_y), fill="#ff7a00", width=3)

        count = len(values)
        step = (right - left) / max(count - 1, 1)
        dot_color = "#f5b400"
        for index, value in enumerate(values):
            x = left + step * index if count > 1 else (left + right) / 2
            y = cls._chart_y(value, top, bottom)
            draw.ellipse((x - 5, y - 5, x + 5, y + 5), fill=dot_color, outline=dot_color)

        label_step = max(1, count // 18)
        for index in range(0, count, label_step):
            x = left + step * index if count > 1 else (left + right) / 2
            text = str(index + 1)
            text_width, _ = cls._text_size(draw, text, small_font)
            draw.text((x - text_width / 2, bottom + 10), text, font=small_font, fill="#444444")

        cls._draw_centered_text(draw, (left + right) / 2, height - 28, "学号", label_font)
        cls._draw_rotated_text(image, "达成度值", label_font, 24, (top + bottom) / 2 - 50)

        legend_x, legend_y = 1170, 235
        draw.line((legend_x, legend_y, legend_x + 32, legend_y), fill="#3f9bd8", width=4)
        draw.text((legend_x + 44, legend_y - 13), "平均值", font=label_font, fill="#333333")
        draw.line((legend_x, legend_y + 42, legend_x + 32, legend_y + 42), fill="#ff7a00", width=4)
        draw.text((legend_x + 44, legend_y + 29), "期望值", font=label_font, fill="#333333")
        draw.ellipse((legend_x + 10, legend_y + 76, legend_x + 22, legend_y + 88), fill=dot_color, outline=dot_color)
        draw.text((legend_x + 44, legend_y + 68), "达成度", font=label_font, fill="#333333")

        image.save(output_path)
        return output_path

    @classmethod
    def _chapter_five_chart_paths(cls, context):
        summary = context.get("summary", {})
        course = summary.get("course")
        course_id = getattr(course, "id", "course")
        chart_dir = cls._chart_output_dir()

        chart_paths = {"overall": None, "objectives": {}}
        overall_path = chart_dir / f"course_{course_id}_chapter5_overall.png"
        chart_paths["overall"] = cls._save_overall_attainment_chart(summary, overall_path)

        for index, objective in enumerate(summary.get("objective_results", []), start=1):
            objective_path = chart_dir / f"course_{course_id}_chapter5_objective_{index}.png"
            saved_path = cls._save_objective_attainment_chart(objective, objective_path, index)
            if saved_path:
                chart_paths["objectives"][objective.get("objective_title")] = saved_path
        return chart_paths

    @classmethod
    def _add_chart_picture(cls, document, image_path, width_cm=15.0):
        if not image_path or not Path(image_path).exists():
            return None
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Cm(width_cm))
        return paragraph

    @staticmethod
    def _format_number(value, digits=2):
        value = float(value or 0)
        if value.is_integer():
            return str(int(value))
        return f"{value:.{digits}f}".rstrip("0").rstrip(".")

    @staticmethod
    def _format_percent_label(value):
        return str(value or "").replace(".0%", "%").replace(".00%", "%")

    @staticmethod
    def _format_semester_for_table(semester):
        match = re.search(r"(\d{4})[-—](\d{4})学年第(\d+)学期", semester or "")
        if match:
            return f"{match.group(1)}—{match.group(2)}第{match.group(3)}学期"
        return semester or ""

    @staticmethod
    def _infer_grade(course, class_scope):
        source_parts = [class_scope, course.class_names or "", course.major or ""]
        try:
            source_parts.extend(student.class_name or "" for student in course.students)
        except Exception:  # noqa: BLE001
            pass
        source = " ".join(filter(None, source_parts))
        full_year = re.search(r"(?<!\d)(20\d{2})级", source)
        if full_year:
            return f"{full_year.group(1)}级"
        short_year = re.search(r"(?<!\d)(2[0-9])级", source)
        if short_year:
            return f"20{short_year.group(1)}级"
        match = re.search(r"(?<!\d)(2[0-9])\d{2}(?!\d)", source)
        if match:
            return f"20{match.group(1)}级"
        return ""

    @staticmethod
    def _is_noisy_class_label(value):
        parts = [item.strip() for item in re.split(r"[、,，\s]+", str(value or "")) if item.strip()]
        return len(parts) > 12 and sum(1 for item in parts if item.isdigit()) >= 12

    @classmethod
    def _resolve_target_students(cls, course, class_scope):
        if class_scope and class_scope != "全部班级":
            return class_scope
        if cls._is_noisy_class_label(course.class_names):
            return class_scope or ""
        return course.class_names or class_scope or ""

    @staticmethod
    def _paragraphs(cell):
        return cell.paragraphs or [cell.add_paragraph()]

    @staticmethod
    def _set_fixed_table_layout(table):
        table.autofit = False
        table_pr = table._tbl.tblPr
        layout = table_pr.find(qn("w:tblLayout"))
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            table_pr.append(layout)
        layout.set(qn("w:type"), "fixed")

    @classmethod
    def _set_table_column_widths(cls, table, widths_cm):
        cls._set_fixed_table_layout(table)
        cls._set_table_width(table, sum(widths_cm))
        table_grid = table._tbl.tblGrid
        if table_grid is not None:
            for child in list(table_grid):
                table_grid.remove(child)
            for width in widths_cm:
                grid_col = OxmlElement("w:gridCol")
                grid_col.set(qn("w:w"), str(int(Cm(width).twips)))
                table_grid.append(grid_col)
        for row in table.rows:
            for index, width in enumerate(widths_cm):
                if index >= len(row.cells):
                    break
                cell = row.cells[index]
                cell.width = Cm(width)
                cell_pr = cell._tc.get_or_add_tcPr()
                for child in list(cell_pr):
                    if child.tag == qn("w:tcW"):
                        cell_pr.remove(child)
                cell_width = OxmlElement("w:tcW")
                cell_width.set(qn("w:w"), str(int(Cm(width).twips)))
                cell_width.set(qn("w:type"), "dxa")
                cell_pr.append(cell_width)

    @staticmethod
    def _set_table_width(table, width_cm):
        table_pr = table._tbl.tblPr
        table_width = table_pr.find(qn("w:tblW"))
        if table_width is None:
            table_width = OxmlElement("w:tblW")
            table_pr.append(table_width)
        table_width.set(qn("w:w"), str(int(Cm(width_cm).twips)))
        table_width.set(qn("w:type"), "dxa")

    @staticmethod
    def _set_cell_borders(cell, **borders):
        cell_pr = cell._tc.get_or_add_tcPr()
        cell_borders = cell_pr.find(qn("w:tcBorders"))
        if cell_borders is None:
            cell_borders = OxmlElement("w:tcBorders")
            cell_pr.append(cell_borders)
        for edge in ("top", "left", "bottom", "right"):
            edge_data = borders.get(edge, {"val": "nil"})
            tag = qn(f"w:{edge}")
            element = cell_borders.find(tag)
            if element is None:
                element = OxmlElement(f"w:{edge}")
                cell_borders.append(element)
            for key, value in edge_data.items():
                element.set(qn(f"w:{key}"), str(value))

    @staticmethod
    def _set_cell_margins(cell, *, top=80, start=80, bottom=80, end=80):
        cell_pr = cell._tc.get_or_add_tcPr()
        margins = cell_pr.find(qn("w:tcMar"))
        if margins is None:
            margins = OxmlElement("w:tcMar")
            cell_pr.append(margins)
        for side, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
            node = margins.find(qn(f"w:{side}"))
            if node is None:
                node = OxmlElement(f"w:{side}")
                margins.append(node)
            node.set(qn("w:w"), str(value))
            node.set(qn("w:type"), "dxa")

    @classmethod
    def _set_cell_text(
        cls,
        cell,
        text,
        *,
        bold=False,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        font_name="宋体",
        size=10.5,
        line_spacing=1.0,
    ):
        cell.text = "" if text is None else str(text)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cls._paragraphs(cell):
            paragraph.alignment = alignment
            paragraph.paragraph_format.line_spacing = line_spacing
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                cls._set_run_font(run, font_name, size, bold)

    @classmethod
    def _format_table(cls, table):
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)

    @classmethod
    def _set_cover_cell_text(cls, cell, text, *, alignment=WD_ALIGN_PARAGRAPH.CENTER, bottom_border=False):
        cell.text = "" if text is None else str(text)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        borders = {"val": "single", "sz": "8", "space": "0", "color": "000000"} if bottom_border else {"val": "nil"}
        cls._set_cell_borders(cell, top={"val": "nil"}, left={"val": "nil"}, bottom=borders, right={"val": "nil"})
        for paragraph in cls._paragraphs(cell):
            paragraph.alignment = alignment
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                cls._set_run_font(run, "仿宋_GB2312", 16, True)

    @classmethod
    def _set_audit_signature_cell(cls, cell):
        cell.text = ""
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
        cls._set_cell_borders(
            cell,
            top={"val": "single", "sz": "8", "space": "0", "color": "000000"},
            left={"val": "single", "sz": "8", "space": "0", "color": "000000"},
            bottom={"val": "single", "sz": "8", "space": "0", "color": "000000"},
            right={"val": "single", "sz": "8", "space": "0", "color": "000000"},
        )
        paragraph = cls._paragraphs(cell)[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.left_indent = Cm(3.55)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(10)
        tab_stops = paragraph.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(7.2), WD_TAB_ALIGNMENT.LEFT)
        tab_stops.add_tab_stop(Cm(9.75), WD_TAB_ALIGNMENT.LEFT)
        tab_stops.add_tab_stop(Cm(12.15), WD_TAB_ALIGNMENT.LEFT)
        run = paragraph.add_run("签名：")
        cls._set_run_font(run, "仿宋", 12)
        run.add_tab()
        run = paragraph.add_run("年")
        cls._set_run_font(run, "仿宋", 12)
        run.add_tab()
        run = paragraph.add_run("月")
        cls._set_run_font(run, "仿宋", 12)
        run.add_tab()
        run = paragraph.add_run("日")
        cls._set_run_font(run, "仿宋", 12)

    @classmethod
    def _add_cover_logo(cls, document):
        seal_path, wordmark_path = cls._cover_logo_paths()
        if not seal_path or not wordmark_path:
            return
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(30)
        run = paragraph.add_run()
        run.add_picture(str(seal_path), width=Cm(2.45))
        run.add_text("   ")
        run.add_picture(str(wordmark_path), width=Cm(6.35))

    @staticmethod
    def _remove_trailing_rows(table, target_count):
        while len(table.rows) > target_count:
            table._tbl.remove(table.rows[-1]._tr)

    @staticmethod
    def _ensure_rows(table, target_count):
        while len(table.rows) < target_count:
            table.add_row()

    @classmethod
    def _update_basic_info_table(cls, table, course, semester, class_scope, student_count):
        rows = [
            ("课程编号", course.code, "课程名称", course.name, "", ""),
            ("课程性质", course.nature or "", "学分", cls._format_number(course.credits, 1), "学时", cls._format_number(course.hours, 0)),
            ("考核方式", course.assessment_method or "", "开课学期", cls._format_semester_for_table(semester), "课程目标达成度期望值", f"{course.expected_value or 0:.2f}"),
            ("授课对象", cls._resolve_target_students(course, class_scope), "", "", "上课人数", str(student_count)),
        ]
        for row_index, values in enumerate(rows):
            for col_index, value in enumerate(values):
                if row_index == 0 and col_index in {4, 5}:
                    continue
                if row_index == 3 and col_index in {2, 3}:
                    continue
                cls._set_cell_text(table.cell(row_index, col_index), value, bold=col_index % 2 == 0)

    @classmethod
    def _collect_requirement_rows(cls, course):
        rows = []
        objectives = sorted(course.objectives, key=lambda item: item.sequence)
        for objective in objectives:
            mappings = list(objective.requirement_maps)
            for mapping in mappings:
                requirement = mapping.requirement
                rows.append(
                    [
                        cls._requirement_title(requirement),
                        cls._requirement_indicator(requirement),
                        f"{objective.title}：{objective.description}",
                        mapping.support_strength,
                    ]
                )
        return rows

    @classmethod
    def _update_requirement_table(cls, table, course):
        rows = cls._collect_requirement_rows(course)
        if not rows:
            return
        cls._set_table_column_widths(table, [2.2, 5.3, 6.4, 1.4])
        cls._ensure_rows(table, len(rows) + 1)
        cls._remove_trailing_rows(table, len(rows) + 1)
        for row_index, row_data in enumerate(rows, start=1):
            for col_index, value in enumerate(row_data):
                cls._set_cell_text(
                    table.cell(row_index, col_index),
                    value,
                    alignment=WD_ALIGN_PARAGRAPH.LEFT if col_index in {1, 2} else WD_ALIGN_PARAGRAPH.CENTER,
                )

    @classmethod
    def _update_assessment_table(cls, table, course):
        assessments = sorted(course.assessments, key=lambda item: item.sequence)
        objectives = sorted(course.objectives, key=lambda item: item.sequence)
        if len(table.columns) != len(assessments) + 2:
            return
        assessment_width = (15.2 - 2.0 - 1.7) / max(len(assessments), 1)
        cls._set_table_column_widths(table, [2.0] + [assessment_width] * len(assessments) + [1.7])
        required_rows = 2 + len(objectives) + 1
        cls._ensure_rows(table, required_rows)
        cls._remove_trailing_rows(table, required_rows)
        last_col = len(assessments) + 1

        for index, assessment in enumerate(assessments, start=1):
            cls._set_cell_text(table.cell(1, index), cls._split_assessment_label(assessment.name), bold=True)

        totals = {assessment.id: 0.0 for assessment in assessments}
        for row_offset, objective in enumerate(objectives, start=2):
            cls._set_cell_text(table.cell(row_offset, 0), objective.title)
            weight_map = {item.assessment_id: item.weight_score for item in objective.assessment_weights}
            for col_index, assessment in enumerate(assessments, start=1):
                score = float(weight_map.get(assessment.id, 0.0) or 0.0)
                totals[assessment.id] += score
                cls._set_cell_text(table.cell(row_offset, col_index), cls._format_number(score, 1) if score else "")
            cls._set_cell_text(table.cell(row_offset, last_col), cls._format_number(objective.weight, 1))

        total_row = 2 + len(objectives)
        cls._set_cell_text(table.cell(total_row, 0), "合计", bold=True)
        for col_index, assessment in enumerate(assessments, start=1):
            cls._set_cell_text(table.cell(total_row, col_index), cls._format_number(totals[assessment.id], 1))
        cls._set_cell_text(table.cell(total_row, last_col), cls._format_number(sum(item.weight for item in objectives), 1), bold=True)

    @classmethod
    def _update_quantitative_table(cls, table, chapter_four):
        rows = chapter_four.get("quantitative_rows") or []
        if not rows:
            return
        cls._set_table_column_widths(table, [2.0, 2.1, 1.5, 1.7, 2.1, 2.0, 2.0])
        cls._ensure_rows(table, len(rows) + 1)
        cls._remove_trailing_rows(table, len(rows) + 1)
        for row_index, row_data in enumerate(rows, start=1):
            values = [
                row_data["objective_title"],
                row_data["assessment_name"],
                cls._format_percent_label(row_data["percentage_label"]),
                cls._format_number(row_data["target_score"], 2),
                f"{row_data['actual_average_score']:.1f}",
                cls._format_number(row_data["objective_weight"], 1),
                f"{row_data['objective_attainment']:.2f}",
            ]
            for col_index, value in enumerate(values):
                cls._set_cell_text(table.cell(row_index, col_index), value)

    @classmethod
    def _update_qualitative_table(cls, table, chapter_four):
        rows = chapter_four.get("qualitative_rows") or []
        if not rows:
            return
        cls._set_table_column_widths(table, [1.8, 6.6, 0.9, 0.9, 0.9, 0.9, 1.3, 1.3])
        cls._ensure_rows(table, len(rows) + 1)
        cls._remove_trailing_rows(table, len(rows) + 1)
        for cell in table.rows[0].cells:
            cls._set_cell_margins(cell, top=50, start=50, bottom=50, end=50)
            cls._set_cell_text(cell, cell.text, bold=True, size=9.5, line_spacing=0.95)
        for row_index, row_data in enumerate(rows, start=1):
            row = table.rows[row_index].cells
            for cell in row:
                cls._set_cell_margins(cell, top=40, start=45, bottom=40, end=45)
            if row_data["objective_title"] == "课程总目标计算":
                cls._set_cell_text(row[0], "课程总目标计算", size=9.5, line_spacing=0.95)
                cls._set_cell_text(row[7], f"{row_data['attainment']:.2f}", size=9.5, line_spacing=0.95)
                continue
            values = [
                row_data["objective_title"],
                row_data["question_ability"],
                row_data["excellent_count"],
                row_data["good_count"],
                row_data["medium_count"],
                row_data["poor_count"],
                f"{row_data['score_rate']:.1f}",
                f"{row_data['attainment']:.2f}",
            ]
            for col_index, value in enumerate(values):
                cls._set_cell_text(
                    row[col_index],
                    value,
                    alignment=WD_ALIGN_PARAGRAPH.LEFT if col_index == 1 else WD_ALIGN_PARAGRAPH.CENTER,
                    size=8.5 if col_index == 1 else 9.5,
                    line_spacing=0.92 if col_index == 1 else 0.95,
                )

    @classmethod
    def _update_qualitative_formula(cls, document, formula):
        if not formula:
            return
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text.startswith("课程目标达成度=") and "*" in text:
                paragraph.text = formula
                for run in paragraph.runs:
                    cls._set_run_font(run, "仿宋", 16)
                return

    @classmethod
    def _update_qualitative_note(cls, document):
        note = "优、良、中、差分别赋分90、80、70、60，得分率为每个课程目标的加权平均值。"
        for paragraph in document.paragraphs:
            if "优、良、中、差分别赋分" in paragraph.text:
                paragraph.text = note
                for run in paragraph.runs:
                    cls._set_run_font(run, "仿宋", 16)
                return

    @classmethod
    def _apply_to_existing_teacher_template(cls, document, course, semester, class_scope, context):
        summary = context["summary"]
        chapter_four = summary["chapter_four"]
        if len(document.tables) >= 1:
            cls._update_basic_info_table(document.tables[0], course, semester, class_scope, summary["student_count"])
        if len(document.tables) >= 2:
            cls._update_requirement_table(document.tables[1], course)
        if len(document.tables) >= 3:
            cls._update_assessment_table(document.tables[2], course)
        if len(document.tables) >= 4:
            cls._update_quantitative_table(document.tables[3], chapter_four)
        if len(document.tables) >= 5:
            cls._update_qualitative_table(document.tables[4], chapter_four)
        cls._update_qualitative_note(document)
        cls._update_qualitative_formula(document, chapter_four.get("qualitative_formula"))
        cls._remove_existing_tail_sections(document)
        if context.get("chapter_five_ready"):
            cls._build_chapter_five(document, context)
        cls._build_other_section(document, context.get("chapter_five_ready", False))
        cls._build_audit_section(document, context.get("chapter_five_ready", False))

    @classmethod
    def _remove_existing_tail_sections(cls, document):
        start_element = None
        for paragraph in document.paragraphs:
            if paragraph.text.strip().startswith("五、"):
                start_element = paragraph._p
                break
        if start_element is None:
            return

        body = document._body._element
        removing = False
        for child in list(body):
            if child is start_element:
                removing = True
            if removing and child.tag != qn("w:sectPr"):
                body.remove(child)

    @classmethod
    def _build_cover(cls, document, course, semester, class_scope):
        cls._add_text_paragraph(document, line_spacing=1.0, space_after=10)
        cls._add_cover_logo(document)
        cls._add_text_paragraph(document, line_spacing=1.0, space_after=24)
        cls._add_text_paragraph(
            document,
            "课程目标达成度分析及持续改进报告",
            font_name="方正公文小标宋",
            size=26,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            line_spacing=1.25,
        )
        cls._add_text_paragraph(document, line_spacing=1.0, space_after=44)

        cover_rows = [
            ("院系（盖章）", course.department or ""),
            ("专      业", course.major or ""),
            ("年      级", cls._infer_grade(course, class_scope)),
            ("课 程 名 称", course.name),
            ("课程负责人", course.course_owner or ""),
            ("任 课 教 师", course.instructors or course.course_owner or ""),
            ("开 课 学 期", semester),
        ]
        table = document.add_table(rows=len(cover_rows), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        cls._set_table_column_widths(table, [4.35, 9.25])
        for index, (label, value) in enumerate(cover_rows):
            row = table.rows[index]
            row.height = Cm(0.98)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            cls._set_cover_cell_text(row.cells[0], label, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
            cls._set_cover_cell_text(row.cells[1], value, alignment=WD_ALIGN_PARAGRAPH.CENTER, bottom_border=True)
        cls._add_text_paragraph(document, line_spacing=1.0, space_after=0)

    @classmethod
    def _build_basic_info(cls, document, course, semester, class_scope, student_count):
        cls._add_section_title(document, "一、")
        table = document.add_table(rows=4, cols=6)
        cls._format_table(table)
        table.cell(0, 3).merge(table.cell(0, 5))
        table.cell(3, 1).merge(table.cell(3, 3))
        rows = [
            ("课程编号", course.code, "课程名称", course.name, "", ""),
            ("课程性质", course.nature or "", "学分", cls._format_number(course.credits, 1), "学时", cls._format_number(course.hours, 0)),
            ("考核方式", course.assessment_method or "", "开课学期", cls._format_semester_for_table(semester), "课程目标达成度期望值", f"{course.expected_value or 0:.2f}"),
            ("授课对象", cls._resolve_target_students(course, class_scope), "", "", "上课人数", str(student_count)),
        ]
        for row_index, values in enumerate(rows):
            for col_index, value in enumerate(values):
                if row_index == 0 and col_index in {4, 5}:
                    continue
                if row_index == 3 and col_index in {2, 3}:
                    continue
                bold = col_index % 2 == 0
                cls._set_cell_text(table.cell(row_index, col_index), value, bold=bold)

    @staticmethod
    def _requirement_title(requirement):
        code = requirement.code or ""
        title = requirement.title or ""
        if code and title and not title.startswith(code):
            return f"{code}-{title}"
        return title or code

    @staticmethod
    def _requirement_indicator(requirement):
        indicator = requirement.indicator_point or ""
        description = requirement.description or ""
        if indicator and description and not description.startswith(indicator):
            return f"{indicator} {description}"
        return description or indicator

    @classmethod
    def _build_requirement_mapping(cls, document, course):
        cls._add_section_title(document, "二、课程目标与毕业要求指标点的对应关系")
        table = document.add_table(rows=1, cols=4)
        cls._format_table(table)
        cls._set_table_column_widths(table, [2.2, 5.3, 6.4, 1.4])
        headers = ["毕业要求", "毕业要求指标点", "课程目标", "支撑强度"]
        for index, header in enumerate(headers):
            cls._set_cell_text(table.rows[0].cells[index], header, bold=True)

        objectives = sorted(course.objectives, key=lambda item: item.sequence)
        for objective in objectives:
            mappings = list(objective.requirement_maps)
            if not mappings:
                row = table.add_row().cells
                cls._set_cell_text(row[0], "暂未配置")
                cls._set_cell_text(row[1], "暂未配置", alignment=WD_ALIGN_PARAGRAPH.LEFT)
                cls._set_cell_text(row[2], f"{objective.title}：{objective.description}", alignment=WD_ALIGN_PARAGRAPH.LEFT)
                cls._set_cell_text(row[3], "-")
                continue
            for mapping in mappings:
                row = table.add_row().cells
                requirement = mapping.requirement
                cls._set_cell_text(row[0], cls._requirement_title(requirement))
                cls._set_cell_text(row[1], cls._requirement_indicator(requirement), alignment=WD_ALIGN_PARAGRAPH.LEFT)
                cls._set_cell_text(row[2], f"{objective.title}：{objective.description}", alignment=WD_ALIGN_PARAGRAPH.LEFT)
                cls._set_cell_text(row[3], mapping.support_strength)

    @staticmethod
    def _split_assessment_label(name):
        if name in {"课后作业", "随堂测试"}:
            return name[:2] + "\n" + name[2:]
        return name

    @classmethod
    def _build_assessment_mapping(cls, document, course):
        cls._add_section_title(document, "三、支撑指标点的课程目标考核分布")
        cls._add_text_paragraph(document, "（一）课程目标与考核方式对应关系", font_name="黑体", size=16)
        cls._add_text_paragraph(document, "（说明：评价方式包括课堂表现、上机实践、结题报告等）", font_name="仿宋", size=16)

        assessments = sorted(course.assessments, key=lambda item: item.sequence)
        objectives = sorted(course.objectives, key=lambda item: item.sequence)
        table = document.add_table(rows=2 + len(objectives) + 1, cols=len(assessments) + 2)
        cls._format_table(table)
        assessment_width = (15.2 - 2.0 - 1.7) / max(len(assessments), 1)
        cls._set_table_column_widths(table, [2.0] + [assessment_width] * len(assessments) + [1.7])
        last_col = len(assessments) + 1
        table.cell(0, 0).merge(table.cell(1, 0))
        table.cell(0, 1).merge(table.cell(0, last_col - 1))
        table.cell(0, last_col).merge(table.cell(1, last_col))
        cls._set_cell_text(table.cell(0, 0), "课程目标", bold=True)
        cls._set_cell_text(table.cell(0, 1), "考核与评价方式及成绩比例（%）", bold=True)
        cls._set_cell_text(table.cell(0, last_col), "课程目标\n权重", bold=True)
        for index, assessment in enumerate(assessments, start=1):
            cls._set_cell_text(table.cell(1, index), cls._split_assessment_label(assessment.name), bold=True)

        totals = {assessment.id: 0.0 for assessment in assessments}
        for row_offset, objective in enumerate(objectives, start=2):
            cls._set_cell_text(table.cell(row_offset, 0), objective.title)
            weight_map = {item.assessment_id: item.weight_score for item in objective.assessment_weights}
            for col_index, assessment in enumerate(assessments, start=1):
                score = float(weight_map.get(assessment.id, 0.0) or 0.0)
                totals[assessment.id] += score
                cls._set_cell_text(table.cell(row_offset, col_index), cls._format_number(score, 1) if score else "")
            cls._set_cell_text(table.cell(row_offset, last_col), cls._format_number(objective.weight, 1))

        total_row = 2 + len(objectives)
        cls._set_cell_text(table.cell(total_row, 0), "合计", bold=True)
        for col_index, assessment in enumerate(assessments, start=1):
            cls._set_cell_text(table.cell(total_row, col_index), cls._format_number(totals[assessment.id], 1))
        cls._set_cell_text(table.cell(total_row, last_col), cls._format_number(sum(item.weight for item in objectives), 1), bold=True)

    @classmethod
    def _merge_group_cells(cls, table, start_row, end_row, columns, values):
        if end_row <= start_row:
            return
        for col_index, value in zip(columns, values):
            cell = table.cell(start_row, col_index).merge(table.cell(end_row, col_index))
            cls._set_cell_text(cell, value)

    @classmethod
    def _build_chapter_four(cls, document, chapter_four):
        cls._add_section_title(document, "四、课程目标达成度评价")
        cls._add_text_paragraph(document, "（一）课程教学目标达成情况定量评价", font_name="黑体", size=16)
        quantitative_table = document.add_table(rows=1, cols=7)
        cls._format_table(quantitative_table)
        cls._set_table_column_widths(quantitative_table, [2.0, 2.1, 1.5, 1.7, 2.1, 2.0, 2.0])
        quantitative_table.cell(0, 1).merge(quantitative_table.cell(0, 2))
        headers = ["课程目标", "评价方式及比例", "", "目标\n分值", "实际平均分", "分目标权重\n%", "目标达成度"]
        for index, header in enumerate(headers):
            if index == 2:
                continue
            cls._set_cell_text(quantitative_table.cell(0, index), header, bold=True)

        group_start = None
        current_objective = None
        group_values = None
        for row_data in chapter_four["quantitative_rows"]:
            row_index = len(quantitative_table.rows)
            row = quantitative_table.add_row().cells
            objective_title = row_data["objective_title"]
            if current_objective != objective_title:
                if group_start is not None:
                    cls._merge_group_cells(quantitative_table, group_start, row_index - 1, [0, 5, 6], group_values)
                current_objective = objective_title
                group_start = row_index
                group_values = [
                    objective_title,
                    cls._format_number(row_data["objective_weight"], 1),
                    f"{row_data['objective_attainment']:.2f}",
                ]

            cls._set_cell_text(row[0], objective_title)
            cls._set_cell_text(row[1], row_data["assessment_name"])
            cls._set_cell_text(row[2], cls._format_percent_label(row_data["percentage_label"]))
            cls._set_cell_text(row[3], cls._format_number(row_data["target_score"], 2))
            cls._set_cell_text(row[4], f"{row_data['actual_average_score']:.1f}")
            cls._set_cell_text(row[5], cls._format_number(row_data["objective_weight"], 1))
            cls._set_cell_text(row[6], f"{row_data['objective_attainment']:.2f}")
        if group_start is not None:
            cls._merge_group_cells(quantitative_table, group_start, len(quantitative_table.rows) - 1, [0, 5, 6], group_values)

        cls._add_text_paragraph(document, "（二）课程教学目标达成情况定性评价", font_name="黑体", size=16)
        cls._add_text_paragraph(document, "优、良、中、差分别赋分90、80、70、60，得分率为每个课程目标的加权平均值。")
        qualitative_table = document.add_table(rows=1, cols=8)
        cls._format_table(qualitative_table)
        cls._set_table_column_widths(qualitative_table, [1.8, 6.6, 0.9, 0.9, 0.9, 0.9, 1.3, 1.3])
        qualitative_headers = ["课程目标", "问题能力", "优", "良", "中", "差", "得分率", "达成度"]
        for index, header in enumerate(qualitative_headers):
            cls._set_cell_margins(qualitative_table.rows[0].cells[index], top=50, start=50, bottom=50, end=50)
            cls._set_cell_text(qualitative_table.rows[0].cells[index], header, bold=True, size=9.5, line_spacing=0.95)
        for row_data in chapter_four["qualitative_rows"]:
            row = qualitative_table.add_row().cells
            for cell in row:
                cls._set_cell_margins(cell, top=40, start=45, bottom=40, end=45)
            if row_data["objective_title"] == "课程总目标计算":
                merged = row[0].merge(row[6])
                cls._set_cell_text(merged, "课程总目标计算", size=9.5, line_spacing=0.95)
                cls._set_cell_text(row[7], f"{row_data['attainment']:.2f}", size=9.5, line_spacing=0.95)
                continue
            cls._set_cell_text(row[0], row_data["objective_title"], size=9.5, line_spacing=0.95)
            cls._set_cell_text(row[1], row_data["question_ability"], alignment=WD_ALIGN_PARAGRAPH.LEFT, size=8.5, line_spacing=0.92)
            cls._set_cell_text(row[2], row_data["excellent_count"], size=9.5, line_spacing=0.95)
            cls._set_cell_text(row[3], row_data["good_count"], size=9.5, line_spacing=0.95)
            cls._set_cell_text(row[4], row_data["medium_count"], size=9.5, line_spacing=0.95)
            cls._set_cell_text(row[5], row_data["poor_count"], size=9.5, line_spacing=0.95)
            cls._set_cell_text(row[6], f"{row_data['score_rate']:.1f}", size=9.5, line_spacing=0.95)
            cls._set_cell_text(row[7], f"{row_data['attainment']:.2f}", size=9.5, line_spacing=0.95)
        cls._add_text_paragraph(document, chapter_four["qualitative_formula"])

    @classmethod
    def _build_chapter_five(cls, document, context):
        cls._add_section_title(document, "五、评价结果分析与持续改进措施")
        cls._add_text_paragraph(document, "（一）课程目标达成情况分析", font_name="黑体", size=16)
        generated_insight = context["generated_insight"]
        chart_paths = cls._chapter_five_chart_paths(context)
        cls._add_chart_picture(document, chart_paths.get("overall"))
        if generated_insight.get("overview_text"):
            cls._add_text_paragraph(document, generated_insight["overview_text"])
        for item in generated_insight.get("objective_analyses", []):
            cls._add_chart_picture(document, chart_paths.get("objectives", {}).get(item.get("objective_title")))
            cls._add_text_paragraph(document, f"{item['objective_title']}：{item['analysis']}")
            if item.get("evidence"):
                cls._add_text_paragraph(document, "证据：" + "；".join(item["evidence"]))
            if item.get("teaching_suggestion"):
                cls._add_text_paragraph(document, f"后续关注点：{item['teaching_suggestion']}")
        cls._add_text_paragraph(document, "（二）改进措施", font_name="黑体", size=16)
        cls._add_text_paragraph(document, "针对以上分析存在的问题，提出以下几点改进措施：")
        for index, action in enumerate(generated_insight.get("improvement_actions", []), start=1):
            title = action.get("title", "")
            related_objective = action.get("related_objective", "")
            problem = action.get("problem", "")
            expected_effect = action.get("expected_effect", "")
            cls._add_text_paragraph(document, f"（{index}）{title}" if title else f"（{index}）改进措施")
            if related_objective:
                cls._add_text_paragraph(document, f"关联目标：{related_objective}")
            if problem:
                cls._add_text_paragraph(document, f"问题定位：{problem}")
            cls._add_text_paragraph(document, f"改进动作：{action['action']}")
            if expected_effect:
                cls._add_text_paragraph(document, f"预期效果：{expected_effect}")

    @classmethod
    def _build_other_section(cls, document, chapter_five_ready: bool):
        cls._add_section_title(document, "六、其它相关说明" if chapter_five_ready else "五、其它相关说明")
        cls._add_text_paragraph(document, "无。")

    @classmethod
    def _build_audit_section(cls, document, chapter_five_ready: bool):
        cls._add_section_title(document, "七、审核意见" if chapter_five_ready else "六、审核意见")
        table = document.add_table(rows=2, cols=2)
        cls._format_table(table)
        table.autofit = False
        cls._set_table_column_widths(table, [1.35, 13.95])
        labels = ["教研\n室审\n核意\n见", "院系\n领导\n审核\n意见"]
        for row_index, label in enumerate(labels):
            table.rows[row_index].height = Cm(8.1)
            table.rows[row_index].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            cls._set_cell_margins(table.cell(row_index, 0), top=80, start=30, bottom=80, end=30)
            cls._set_cell_margins(table.cell(row_index, 1), top=80, start=120, bottom=120, end=120)
            cls._set_cell_text(table.cell(row_index, 0), label, bold=True, size=12, line_spacing=1.15)
            cls._set_audit_signature_cell(table.cell(row_index, 1))
