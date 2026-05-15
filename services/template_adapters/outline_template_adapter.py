import re
from pathlib import Path

from docx import Document


class OutlineTemplateAdapter:
    """解析课程教学大纲 docx，并输出统一结构。"""

    NAME = "outline_template_adapter"

    FIELD_LABELS = {
        "课程编码": "course_code",
        "课程编号": "course_code",
        "课程名称": "course_name",
        "课程中文名称": "course_name",
        "英文名称": "english_name",
        "课程英文名称": "english_name",
        "课程性质": "nature",
        "课程类别": "category",
        "学时/学分": "hours_credits",
        "总学时": "hours",
        "学分": "credits",
        "适用专业": "major",
        "授课对象": "class_names",
        "授课班级": "class_names",
        "上课班级": "class_names",
        "开课单位": "department",
        "开课时间": "semester_hint",
        "开课学期": "semester_hint",
        "考核形式": "assessment_method",
        "考核方式": "assessment_method",
        "课程负责人": "course_owner",
        "先修课程": "prerequisites",
        "选用教材": "textbook",
    }

    @classmethod
    def extract(cls, file_path: Path):
        document = Document(file_path)
        paragraphs = [item.text.strip() for item in document.paragraphs if item.text.strip()]
        tables = document.tables
        raw_text = "\n".join(paragraphs)

        payload = {
            "course_name": "",
            "course_code": "",
            "english_name": "",
            "nature": "",
            "category": "",
            "hours_credits": "",
            "hours": "",
            "credits": "",
            "major": "",
            "class_names": "",
            "department": "",
            "semester_hint": "",
            "assessment_method": "",
            "course_owner": "",
            "prerequisites": "",
            "textbook": "",
            "description": cls._extract_description(paragraphs, raw_text),
            "objectives": cls._extract_objectives(raw_text),
            "requirements": [],
            "teaching_support": [],
            "assessment_support": [],
            "source_template": file_path.name,
        }

        cls._extract_from_paragraphs(paragraphs, payload)
        cls._extract_from_tables(tables, payload)
        payload["confidence"] = cls._build_confidence(payload)
        return {
            "raw_text": raw_text,
            "payload": payload,
            "summary": "已从教学大纲中提取课程基本信息、课程目标、毕业要求映射、教学内容支撑与考核支撑关系。",
        }

    @classmethod
    def _extract_from_paragraphs(cls, paragraphs, payload):
        for paragraph in paragraphs:
            if "：" not in paragraph and ":" not in paragraph:
                continue
            label, value = re.split(r"[:：]", paragraph, maxsplit=1)
            cls._apply_field(payload, label, value)

        if not payload["course_name"]:
            payload["course_name"] = cls._search(r"课程名称[:：]\s*([^\n]+)", "\n".join(paragraphs))
        if not payload["course_code"]:
            payload["course_code"] = cls._search(r"课程编码[:：]\s*([^\n]+)", "\n".join(paragraphs))

    @classmethod
    def _extract_from_tables(cls, tables, payload):
        for table in tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if not rows:
                continue

            first_row = rows[0]
            first_row_text = "".join(first_row)

            for row in rows:
                for index in range(0, len(row) - 1, 2):
                    cls._apply_field(payload, row[index], row[index + 1])

            if "毕业要求" in first_row_text and "课程目标" in first_row_text:
                header_map = cls._build_header_map(first_row)
                requirement_index = header_map.get("requirement", 0)
                indicator_index = header_map.get("indicator", 1 if len(first_row) > 1 else requirement_index)
                objective_index = header_map.get("objective", 2 if len(first_row) > 2 else indicator_index)
                strength_index = header_map.get("strength")
                for row in rows[1:]:
                    if len(row) <= max(requirement_index, indicator_index, objective_index) or not row[requirement_index]:
                        continue
                    requirement_text = row[requirement_index].replace("\n", " ").strip()
                    indicator_text = row[indicator_index].replace("\n", " ").strip()
                    objective_text = row[objective_index].replace("\n", " ").strip()
                    support_text = row[strength_index].replace("\n", " ").strip() if strength_index is not None and len(row) > strength_index else objective_text
                    indicator_point = cls._search(r"(\d+-\d+)", indicator_text) or indicator_text
                    requirement_code = cls._search(r"^\s*(\d+)", requirement_text) or (indicator_point.split("-", 1)[0] if "-" in indicator_point else indicator_point)
                    objective_ref = cls._normalize_objective_ref(objective_text) or objective_text
                    support_strength = cls._search(r"(H|M|L)", support_text) or "M"
                    payload["requirements"].append(
                        {
                            "requirement_title": requirement_text or requirement_code,
                            "indicator_point": indicator_point,
                            "requirement_description": indicator_text,
                            "objective_ref": objective_ref,
                            "support_strength": support_strength,
                        }
                    )
            elif "教学内容" in first_row_text and ("课程目标" in first_row_text or "教学目标" in first_row_text):
                for row in rows[1:]:
                    if len(row) < 4 or not row[0]:
                        continue
                    payload["teaching_support"].append(
                        {
                            "sequence": row[0].strip(),
                            "content": row[1].replace("\n", " ").strip(),
                            "teaching_mode": row[3].replace("\n", " ").strip() if len(row) > 3 else "",
                            "objective_ref": cls._normalize_objective_ref(row[-1]) or row[-1].replace("\n", " ").strip(),
                        }
                    )
            elif "课程目标" in first_row_text and "权重" in first_row_text:
                header = rows[1] if len(rows) > 1 and "课程目标" in rows[1][0] else first_row
                assessment_names = [item.replace("\n", "").strip() for item in header[1:-1]]
                data_rows = rows[2:] if header is not first_row else rows[1:]
                for row in data_rows:
                    if len(row) < 3 or not row[0]:
                        continue
                    if "合计" in row[0]:
                        continue
                    mapping = []
                    for name, value in zip(assessment_names, row[1:-1]):
                        if name and str(value).strip():
                            mapping.append({"assessment_name": name, "weight_score": str(value).strip()})
                    payload["assessment_support"].append(
                        {
                            "objective_ref": cls._normalize_objective_ref(row[0]) or row[0].strip(),
                            "objective_weight": row[-1].strip(),
                            "assessment_map": mapping,
                        }
                    )

    @classmethod
    def _apply_field(cls, payload, raw_label: str, raw_value: str):
        label = cls._normalize_label(raw_label)
        value = str(raw_value or "").strip()
        if not label or not value:
            return
        normalized_map = {cls._normalize_label(key): target for key, target in cls.FIELD_LABELS.items()}
        key = normalized_map.get(label)
        if not key:
            return
        if key == "category":
            hours_match = re.search(r"(.+?)\s*学时/学分\s*[:：]\s*([0-9.]+)\s*/\s*([0-9.]+)", value)
            if hours_match:
                payload["category"] = hours_match.group(1).strip()
                payload["hours"] = hours_match.group(2).strip()
                payload["credits"] = hours_match.group(3).strip()
                payload["hours_credits"] = f"{payload['hours']}/{payload['credits']}"
                return
        payload[key] = value
        if key == "hours_credits" and "/" in value:
            hours, credits = value.split("/", 1)
            payload["hours"] = hours.strip()
            payload["credits"] = credits.strip()
        if key in {"hours", "credits"} and payload.get("hours") and payload.get("credits"):
            payload["hours_credits"] = f"{payload['hours']}/{payload['credits']}"

    @staticmethod
    def _normalize_label(value: str) -> str:
        return re.sub(r"\s+", "", str(value or "").strip())

    @classmethod
    def _build_header_map(cls, header_row):
        header_map = {}
        for index, header in enumerate(header_row):
            label = cls._normalize_label(header)
            if not label:
                continue
            if label == "毕业要求":
                header_map["requirement"] = index
            elif "毕业要求指标点" in label or "指标点" in label:
                header_map["indicator"] = index
            elif "课程目标" in label:
                header_map["objective"] = index
            elif "支撑强度" in label:
                header_map["strength"] = index
        return header_map

    @staticmethod
    def _chinese_digit_to_int(value: str):
        text = str(value or "").strip()
        if not text:
            return None
        if text.isdigit():
            return int(text)
        number_map = {
            "一": 1,
            "二": 2,
            "三": 3,
            "四": 4,
            "五": 5,
            "六": 6,
            "七": 7,
            "八": 8,
            "九": 9,
            "十": 10,
        }
        return number_map.get(text)

    @classmethod
    def _normalize_objective_ref(cls, value: str) -> str:
        text = str(value or "").replace("\n", " ").strip()
        match = re.search(r"课程目标\s*([0-9一二三四五六七八九十]+)", text)
        if not match:
            return ""
        number = cls._chinese_digit_to_int(match.group(1))
        return f"课程目标{number}" if number else ""

    @staticmethod
    def _extract_block(raw_text: str, start_marker: str, end_marker: str) -> str:
        match = re.search(rf"{re.escape(start_marker)}\s*(.+?)\s*{re.escape(end_marker)}", raw_text, re.S)
        return match.group(1).strip() if match else ""

    @classmethod
    def _extract_description(cls, paragraphs, raw_text: str) -> str:
        numbered = cls._extract_block(raw_text, "二、课程简介", "三、课程目标")
        if numbered:
            return numbered

        collected = []
        in_description = False
        for paragraph in paragraphs:
            text = paragraph.strip()
            normalized = cls._normalize_label(text)
            if normalized == "课程简介":
                in_description = True
                continue
            if in_description and normalized in {"课程目标", "一、课程目标", "（一）课程目标"}:
                break
            if in_description and text:
                collected.append(text)
        return "\n".join(collected).strip()

    @classmethod
    def _extract_objectives(cls, raw_text: str):
        objectives = []
        seen = set()
        for line in raw_text.splitlines():
            match = re.match(r"^\s*课程目标\s*([0-9一二三四五六七八九十]+)\s*[:：]\s*(.+?)\s*$", line)
            if not match:
                continue
            number = cls._chinese_digit_to_int(match.group(1))
            description = match.group(2).strip()
            if not number or not description:
                continue
            title = f"课程目标{number}"
            if title in seen:
                continue
            seen.add(title)
            objectives.append({"title": title, "description": description})
        return objectives

    @staticmethod
    def _search(pattern: str, text: str) -> str:
        match = re.search(pattern, text, re.S)
        return match.group(1).strip() if match else ""

    @staticmethod
    def _build_confidence(payload) -> float:
        score = 0.0
        for key in ("course_name", "course_code", "hours_credits", "assessment_method", "description"):
            if payload.get(key):
                score += 0.12
        if payload.get("objectives"):
            score += 0.2
        if payload.get("requirements"):
            score += 0.16
        if payload.get("assessment_support"):
            score += 0.16
        return round(min(score, 0.98), 2)
